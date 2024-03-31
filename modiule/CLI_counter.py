import os
import re
from collections import Counter
from PyPDF2 import PdfFileReader, PdfFileWriter
from docx import Document

def count_letters(file_path):
    if file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    elif file_path.endswith('.pdf'):
        with open(file_path, 'rb') as file:
            pdf_reader = PdfFileReader(file)
            text = ''
            for page_num in range(pdf_reader.numPages):
                text += pdf_reader.getPage(page_num).extractText()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        text = ''
        for para in doc.paragraphs:
            text += para.text
    
    # Remove non-alphabetic characters and convert text to lowercase
    text = re.sub(r'[^a-zA-Z]', '', text.lower())
    
    # Count the occurrences of each letter
    letter_counts = Counter(text)
    
    # Calculate the total number of letters
    total_letters = sum(letter_counts.values())
    
    # Calculate the percentage of each letter
    letter_percentages = {letter: (count / total_letters) * 100 for letter, count in letter_counts.items()}
    
    return letter_counts, letter_percentages

def save_results(file_name, letter_counts, letter_percentages, file_format):
    if file_format == 'txt':
        with open(file_name + '.txt', 'w') as txt_file:
            txt_file.write("Letter counts:\n")
            for letter, count in sorted(letter_counts.items()):
                txt_file.write(f"{letter}: {count}\n")
            txt_file.write("\nLetter percentages:\n")
            for letter, percentage in sorted(letter_percentages.items()):
                txt_file.write(f"{letter}: {percentage:.2f}%\n")
    elif file_format == 'docx':
        doc = Document()
        doc.add_heading('Letter Counts and Percentages', level=1)
        doc.add_paragraph("Letter counts:")
        for letter, count in sorted(letter_counts.items()):
            doc.add_paragraph(f"{letter}: {count}")
        doc.add_paragraph("\nLetter percentages:")
        for letter, percentage in sorted(letter_percentages.items()):
            doc.add_paragraph(f"{letter}: {percentage:.2f}%")
        doc.save(file_name + '.docx')
    elif file_format == 'pdf':
        pdf_writer = PdfFileWriter()
        pdf_writer.add_blank_page()  # Add a blank page to start the PDF
        pdf = open(file_name + '.pdf', 'wb')
        pdf_writer.write(pdf)
        pdf.close()
        pdf = open(file_name + '.pdf', 'rb+')
        pdf_reader = PdfFileReader(pdf)
        page = pdf_reader.getPage(0)

        text_content = f"Letter counts:\n"
        for letter, count in sorted(letter_counts.items()):
            text_content += f"{letter}: {count}\n"
        text_content += f"\nLetter percentages:\n"
        for letter, percentage in sorted(letter_percentages.items()):
            text_content += f"{letter}: {percentage:.2f}%\n"

        page.mergePage(PdfFileReader(open(file_name + '.txt', 'rb')).getPage(0))
        pdf_writer.write(pdf)
        pdf.close()
        os.remove(file_name + '.txt')  # Remove the temporary text file

def main():
    file_path = input("Enter the path of the file (.txt, .pdf, or .docx): ")
    if os.path.exists(file_path):
        letter_counts, letter_percentages = count_letters(file_path)
        file_name = os.path.splitext(os.path.basename(file_path))[0] + "_results"
        
        print("Select output format:")
        print("1. Text (.txt)")
        print("2. Microsoft Word (.docx)")
        print("3. PDF (.pdf)")
        choice = input("Enter your choice: ")
        
        if choice == '1':
            save_results(file_name, letter_counts, letter_percentages, 'txt')
        elif choice == '2':
            save_results(file_name, letter_counts, letter_percentages, 'docx')
        elif choice == '3':
            save_results(file_name, letter_counts, letter_percentages, 'pdf')
        else:
            print("Invalid choice. Exiting...")
            return
        
        print("Results saved successfully.")
    else:
        print("File not found.")

if __name__ == "__main__":
    main()
